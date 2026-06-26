"""
将 产品说明.md 转换为 A4 Word 文档，保留 Markdown 结构与中文排版。
字体：正文思源宋体(Noto Serif SC) / 标题微软雅黑，西文 Times New Roman。
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ========== 配置 ==========
INPUT_MD = os.path.join(os.path.dirname(__file__), '产品说明.md')
OUTPUT_DOCX = os.path.join(os.path.dirname(__file__), '《卅一景》产品说明.docx')

FONT_BODY_CN = '思源宋体'       # 正文中文
FONT_BODY_EN = 'Times New Roman' # 正文西文
FONT_HEADING_CN = '微软雅黑'     # 标题中文
FONT_HEADING_EN = 'Arial'        # 标题西文
FONT_CODE = 'Consolas'           # 代码块
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_CODE = Pt(9)
LINE_SPACING = 1.5

# ========== 工具函数 ==========

def set_run_font(run, cn_font, en_font, size, bold=False, color=None):
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    if color:
        run.font.color.rgb = color

def set_paragraph_spacing(para, before=0, after=0, line_spacing=LINE_SPACING):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line_spacing

def add_heading_styled(doc, text, level):
    para = doc.add_paragraph()
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = FONT_SIZE_H1
        set_paragraph_spacing(para, before=24, after=12)
    elif level == 2:
        size = FONT_SIZE_H2
        set_paragraph_spacing(para, before=18, after=8)
    else:
        size = FONT_SIZE_H3
        set_paragraph_spacing(para, before=12, after=6)
    run = para.add_run(text)
    set_run_font(run, FONT_HEADING_CN, FONT_HEADING_EN, size, bold=True)
    return para

def add_body_paragraph(doc, text, indent=False):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=2, after=2)
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)
    # 处理加粗 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, FONT_BODY_CN, FONT_BODY_EN, FONT_SIZE_BODY, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, FONT_BODY_CN, FONT_BODY_EN, FONT_SIZE_BODY)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=1, after=1)
    para.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    para.paragraph_format.first_line_indent = Cm(-0.4)
    bullet_char = '•' if level == 0 else '◦'
    parts = re.split(r'(\*\*.*?\*\*)', text)
    first = True
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, FONT_BODY_CN, FONT_BODY_EN, FONT_SIZE_BODY, bold=True)
        else:
            if first:
                run = para.add_run(f'{bullet_char} {part}')
                first = False
            else:
                run = para.add_run(part)
            set_run_font(run, FONT_BODY_CN, FONT_BODY_EN, FONT_SIZE_BODY)
    return para

def add_code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        set_paragraph_spacing(para, before=0, after=0, line_spacing=1.2)
        para.paragraph_format.left_indent = Cm(1.0)
        run = para.add_run(line)
        set_run_font(run, FONT_CODE, FONT_CODE, FONT_SIZE_CODE)
        # 灰色背景效果用浅灰字色模拟
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def set_cell_font(cell, text, bold=False):
    cell.text = ''
    para = cell.paragraphs[0]
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, FONT_BODY_CN, FONT_BODY_EN, Pt(10), bold=True)
        else:
            # 处理 `code` 内联
            code_parts = re.split(r'(`.*?`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = para.add_run(cp[1:-1])
                    set_run_font(run, FONT_CODE, FONT_CODE, Pt(9.5))
                else:
                    run = para.add_run(cp)
                    set_run_font(run, FONT_BODY_CN, FONT_BODY_EN, Pt(10), bold=bold)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True)
    # 内容
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            set_cell_font(table.rows[r_idx + 1].cells[c_idx], cell_text)
    # 简单边框样式
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '999999')
        borders.append(el)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    doc.add_paragraph()  # 表后空行

# ========== 层级编号 ==========

CN_NUMBERS = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
              '十一', '十二', '十三', '十四', '十五']

# ========== 解析 Markdown ==========

def parse_md(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    return [l.rstrip('\n') for l in lines]

def generate_docx(md_lines):
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    h2_counter = 0   # ## 级别计数
    h3_counter = 0   # ### 级别计数

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(md_lines):
        line = md_lines[i]

        # 代码块
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if not in_table:
                # 可能是表头
                # 检查下一行是否是分隔线
                if i + 1 < len(md_lines) and re.match(r'\|[\s\-:|]+\|', md_lines[i+1].strip()):
                    in_table = True
                    table_headers = cells
                    table_rows = []
                    i += 2  # 跳过分隔行
                    continue
                else:
                    in_table = True
                    table_headers = cells
                    table_rows = []
                    i += 1
                    continue
            else:
                table_rows.append(cells)
                # 检查下一行是否还是表格
                if i + 1 >= len(md_lines) or not md_lines[i+1].strip().startswith('|'):
                    add_table(doc, table_headers, table_rows)
                    in_table = False
                i += 1
                continue
        elif in_table:
            add_table(doc, table_headers, table_rows)
            in_table = False
            # 不增加 i，继续处理当前行
            continue

        # 分隔线
        if line.strip() == '---':
            i += 1
            continue

        # 空行
        if line.strip() == '':
            i += 1
            continue

        # 标题
        if line.startswith('# '):
            add_heading_styled(doc, line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith('## '):
            h2_counter += 1
            h3_counter = 0  # 重置子序号
            title = f'{CN_NUMBERS[h2_counter - 1]}、{line[3:].strip()}'
            add_heading_styled(doc, title, 2)
            i += 1
            continue
        if line.startswith('### '):
            h3_counter += 1
            title = f'{h2_counter}.{h3_counter} {line[4:].strip()}'
            add_heading_styled(doc, title, 3)
            i += 1
            continue

        # 列表项
        if line.strip().startswith('- '):
            text = line.strip()[2:]
            level = (len(line) - len(line.lstrip())) // 2
            add_bullet(doc, text, level)
            i += 1
            continue

        # 普通段落
        add_body_paragraph(doc, line.strip())
        i += 1

    # 如果最后还有未关闭的表格
    if in_table:
        add_table(doc, table_headers, table_rows)

    doc.save(OUTPUT_DOCX)
    print(f'已生成: {OUTPUT_DOCX}')

if __name__ == '__main__':
    md_lines = parse_md(INPUT_MD)
    generate_docx(md_lines)
